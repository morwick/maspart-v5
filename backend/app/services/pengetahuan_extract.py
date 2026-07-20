"""Ekstraksi isi berkas → "bagian" mentah untuk diindeks jadi chunk.

Satu bagian = {teks, tabel, halaman, sub, gambar[]} di mana `gambar` = list bytes
PNG/JPG yang ditemukan di bagian itu. Modul ini SENGAJA bodoh: ia hanya membaca &
memotong, tidak menyentuh store, tidak memanggil LLM.

Semua parser berat di-LAZY IMPORT: server 3,8 GB dengan torch/DINOv2 ~1 GB sudah
sesak, jadi RSS idle tak boleh naik hanya karena fitur ini terpasang. Guard ukuran
& jumlah halaman/sheet/gambar ada di sini, bukan di pemanggil.

Pakai ulang: `ai_sheet._read_csv` (decode utf-8-sig→cp1252 + sniffer delimiter)
dan konstanta MAX_COLS/MAX_ROWS-nya, supaya perilaku CSV sama persis dengan
fitur unggah Excel ke chat yang sudah dipakai.
"""
from __future__ import annotations

import hashlib
import io
import re
import zipfile
from pathlib import Path

from . import ai_sheet

# ── plafon (RAM & waktu) ─────────────────────────────────────────────
MAX_TXT_BYTES = 2 * 1024 * 1024
MAX_SHEET = 12
MAX_BARIS_TABEL = 30          # baris per chunk tabel
MAX_KOLOM_TABEL = 12
MAX_GAMBAR = 60
MAX_HALAMAN_PDF = 200
MIN_SISI_GAMBAR = 120            # px — di bawah ini biasanya ikon/garis hias
MAX_PX_GAMBAR = 30_000_000       # ditolak SEBELUM decode (kunci RAM)
MAX_GAMBAR_HALAMAN = 4
MAX_GAMBAR_BYTES = 48 * 1024 * 1024   # total PNG teroptimasi ditahan di RAM
LOGO_MIN_HALAMAN = 3             # hash sama di ≥3 halaman → kop/watermark
MAX_CAPTION = 300
# .docx/.xlsx = arsip ZIP — file kecil bisa mengembang jadi ratusan MB saat
# di-ekstrak (zip-bomb). Cek ukuran TERURAI sebelum membuka isinya.
MAX_ZIP_URAI = 100 * 1024 * 1024
MAX_ZIP_ENTRI = 2000

EKSTENSI = (".pdf", ".xlsx", ".xlsm", ".csv", ".docx", ".txt",
            ".png", ".jpg", ".jpeg")

# ── chunking ─────────────────────────────────────────────────────────
TARGET_CHUNK = 1200
MAX_CHUNK = 1800
OVERLAP = 120


def _potong_di_batas(s: str, batas: int) -> int:
    """Titik potong terbaik ≤ batas: paragraf > baris > kalimat > batas kata."""
    for pola in ("\n\n", "\n", ". "):
        i = s.rfind(pola, batas // 2, batas)
        if i > 0:
            return i + len(pola)
    i = s.rfind(" ", batas // 2, batas)
    return i + 1 if i > 0 else batas


def potong_teks(teks: str, target: int = TARGET_CHUNK) -> list[str]:
    """Pecah prosa jadi potongan ~target char dengan overlap kecil, tak pernah
    memotong di tengah kata. Overlap menjaga kalimat yang terbelah tetap utuh
    di salah satu potongan."""
    s = (teks or "").strip()
    if not s:
        return []
    if len(s) <= MAX_CHUNK:
        return [s]
    out: list[str] = []
    pos = 0
    while pos < len(s):
        sisa = s[pos:]
        if len(sisa) <= MAX_CHUNK:
            out.append(sisa.strip())
            break
        cut = _potong_di_batas(sisa, target)
        out.append(sisa[:cut].strip())
        pos += max(cut - OVERLAP, 1)
    return [c for c in out if c]


def _bagian(teks="", tabel=None, halaman=0, sub="", gambar=None,
            jalur=None, kolom=None, gambar_meta=None) -> dict:
    """Satu bagian mentah. Field baru semuanya OPSIONAL supaya pemanggil lama
    tetap sah: `jalur` = breadcrumb bab→subbab, `kolom` = nama kolom tabel,
    `gambar_meta` = keterangan paralel `gambar` ({caption, halaman})."""
    return {"teks": teks or "", "tabel": tabel or [], "halaman": halaman,
            "sub": sub or "", "gambar": gambar or [],
            "jalur": list(jalur or []), "kolom": list(kolom or []),
            "gambar_meta": list(gambar_meta or [])}


def _rapikan_baris(baris) -> list[str]:
    out = []
    for c in list(baris)[:MAX_KOLOM_TABEL]:
        out.append("" if c is None else str(c).strip())
    while out and not out[-1]:
        out.pop()
    return out


def _tabel_jadi_bagian(rows: list[list], halaman: int, sub: str,
                       jalur=None) -> list[dict]:
    """Baris tabel → bagian ber-`tabel`, dipotong per MAX_BARIS_TABEL dengan
    header diulang di tiap potongan supaya tiap chunk berdiri sendiri.

    `kolom` (nama header) disimpan terpisah agar tabel bisa DITEMUKAN lewat
    pertanyaan tentang kolomnya ("tabel mana yang punya masa garansi?").
    """
    bersih = [_rapikan_baris(r) for r in rows]
    bersih = [r for r in bersih if any(r)]
    if not bersih:
        return []
    header = bersih[0]
    kolom = [c for c in header if c][:MAX_KOLOM_TABEL]
    isi = bersih[1:] or []
    if not isi:
        return [_bagian(tabel=[header], halaman=halaman, sub=sub,
                        jalur=jalur, kolom=kolom)]
    out = []
    for i in range(0, len(isi), MAX_BARIS_TABEL):
        potong = isi[i:i + MAX_BARIS_TABEL]
        b = _bagian(tabel=[header, *potong],
                    halaman=halaman + i if halaman else i + 2, sub=sub,
                    jalur=jalur, kolom=kolom)
        b["baris_total"] = len(isi)
        b["baris_dari"] = i
        out.append(b)
    return out


# ── format sederhana ─────────────────────────────────────────────────
def dari_txt(data: bytes) -> list[dict]:
    if len(data) > MAX_TXT_BYTES:
        data = data[:MAX_TXT_BYTES]
    teks = None
    for enc in ("utf-8-sig", "cp1252"):
        try:
            teks = data.decode(enc)
            break
        except Exception:
            teks = None
    if teks is None:
        raise ValueError("Berkas teks tak terbaca (encoding tak dikenal).")
    return [_bagian(teks=t) for t in potong_teks(teks)]


def dari_csv(data: bytes) -> list[dict]:
    rows = ai_sheet._read_csv(data)
    if rows is None:
        raise ValueError("Berkas CSV tak terbaca (encoding/format tak dikenal).")
    return _tabel_jadi_bagian(rows, 0, "")


def _cek_zip(data: bytes) -> None:
    """Guard zip-bomb untuk .xlsx/.docx — cek ukuran TERURAI sebelum parse."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            infos = z.infolist()
            if len(infos) > MAX_ZIP_ENTRI:
                raise ValueError("Berkas ditolak: terlalu banyak isi di dalamnya.")
            if sum(i.file_size for i in infos) > MAX_ZIP_URAI:
                raise ValueError("Berkas ditolak: isinya mengembang terlalu besar.")
    except zipfile.BadZipFile:
        raise ValueError("Berkas rusak / bukan format yang valid.")


MAX_GAMBAR_EXCEL = 20
_EXCEL_GAMBAR_MAKS_BYTES = 6 * 1024 * 1024   # pass kedua memuat workbook PENUH


def _gambar_excel(data: bytes, out: list[dict]) -> None:
    """Tempelkan gambar embedded Excel ke bagian tabel yang memuat barisnya.

    `read_only=True` TIDAK mengisi `ws._images` — itu sebabnya V1 buta terhadap
    gambar Excel. Pass kedua ini memuat workbook penuh, jadi dibatasi ukuran
    berkas dan seluruhnya best-effort: gagal = tak ada gambar, bukan gagal
    indexing. `_images`/`_data()` API privat openpyxl — kalau hilang di versi
    mendatang, fiturnya mati sendiri.
    """
    if len(data) > _EXCEL_GAMBAR_MAKS_BYTES:
        return
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), data_only=True)
    except Exception:
        return
    n = 0
    try:
        for nama in list(wb.sheetnames)[:MAX_SHEET]:
            try:
                ws = wb[nama]
                for img in list(getattr(ws, "_images", []) or []):
                    if n >= MAX_GAMBAR_EXCEL:
                        return
                    try:
                        baris = int(img.anchor._from.row)
                    except Exception:
                        baris = 0
                    try:
                        png = optimalkan_gambar(img._data())
                    except Exception:
                        continue
                    # Bagian sheet yang rentang barisnya memuat anchor gambar;
                    # di luar rentang → bagian pertama sheet itu.
                    kandidat = [b for b in out if b.get("sub") == nama]
                    if not kandidat:
                        continue
                    target = next(
                        (b for b in kandidat
                         if b.get("baris_dari", 0) <= baris
                         < b.get("baris_dari", 0) + MAX_BARIS_TABEL),
                        kandidat[0])
                    cap = _caption([nama], "",
                                   " ".join(str(c) for c in (target.get("tabel") or [[]])[0]))
                    target["gambar"].append(png)
                    target["gambar_meta"].append({"caption": cap, "halaman": 0})
                    n += 1
            except Exception:
                continue
    finally:
        try:
            wb.close()
        except Exception:
            pass


def dari_excel(data: bytes) -> list[dict]:
    """Semua sheet (maks MAX_SHEET) → bagian tabel. `read_only=True` supaya
    openpyxl streaming, tak memuat seluruh workbook ke RAM."""
    _cek_zip(data)
    from openpyxl import load_workbook
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception:
        raise ValueError("Berkas bukan Excel yang valid / rusak.")
    out: list[dict] = []
    try:
        for nama in list(wb.sheetnames)[:MAX_SHEET]:
            try:
                ws = wb[nama]
                rows: list[list] = []
                for r in ws.iter_rows(values_only=True):
                    rows.append(list(r[:ai_sheet.MAX_COLS]))
                    if len(rows) > ai_sheet.MAX_ROWS:
                        break
                out.extend(_tabel_jadi_bagian(rows, 0, nama, [nama]))
            except Exception:
                continue    # satu sheet rusak tak boleh menjatuhkan sisanya
    finally:
        wb.close()
    _gambar_excel(data, out)
    return out


def dari_gambar(data: bytes) -> list[dict]:
    return [_bagian(gambar=[optimalkan_gambar(data)])]


# ── gambar ───────────────────────────────────────────────────────────
def optimalkan_gambar(data: bytes, maks_sisi: int = 1600) -> bytes:
    """Normalkan ke PNG, perkecil bila raksasa. `MAX_IMAGE_PIXELS` dibatasi
    supaya gambar decompression-bomb ditolak Pillow, bukan menghabiskan RAM."""
    from PIL import Image
    Image.MAX_IMAGE_PIXELS = 64_000_000
    try:
        im = Image.open(io.BytesIO(data))
        im.load()
    except Exception:
        raise ValueError("Gambar rusak / format tak dikenal.")
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    if max(im.size) > maks_sisi:
        rasio = maks_sisi / max(im.size)
        im = im.resize((max(int(im.width * rasio), 1), max(int(im.height * rasio), 1)))
    buf = io.BytesIO()
    im.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


# ── PDF (lazy import pypdfium2) ──────────────────────────────────────
_SPASI2 = re.compile(r"\s{2,}")


def _baris_tabel_pdf(teks: str) -> list[list[str]]:
    """Rekonstruksi tabel dari tata letak whitespace: baris dengan ≥3 pemisah
    '2+ spasi' dianggap baris tabel. Ini heuristik — `catatan` tool memberi tahu
    model agar tidak mengklaim presisi kolom. PDF yang tabelnya penting sebaiknya
    diunggah admin dalam bentuk Excel/CSV (jalur akurat)."""
    baris = []
    for ln in (teks or "").splitlines():
        # 3 kolom = 2 pemisah; ambang di JUMLAH KOLOM, bukan jumlah pemisah
        if len(_SPASI2.findall(ln)) >= 2:
            sel = [c.strip() for c in _SPASI2.split(ln.strip()) if c.strip()]
            if len(sel) >= 3:
                baris.append(sel[:MAX_KOLOM_TABEL])
                continue
        baris.append(None)   # penanda putus
    out, blok = [], []
    for b in baris:
        if b is None:
            if len(blok) >= 2:
                out.append(blok)
            blok = []
        else:
            blok.append(b)
    if len(blok) >= 2:
        out.append(blok)
    return out[0] if len(out) == 1 else [r for blok in out for r in blok]


_FIGUR_RE = re.compile(
    r"^\s*(?:gambar|gbr|fig(?:ure)?|图表|图|表)\s*[\d.\-]*\s*[:.．、]?\s*(?P<isi>\S.*)$",
    re.I)
# Judul bernomor (1.2.3 Judul) & bab Mandarin — dipakai bila PDF tak punya outline.
_HEAD_NOMOR_RE = re.compile(r"^\s*(\d+(?:\.\d+){0,3})[\s.)、]+(\S.{0,80})$")
_HEAD_CN_RE = re.compile(r"^\s*(第[一二三四五六七八九十百\d]+[章节節])\s*(.{0,60})$")


def _baris_figur(teks: str) -> str:
    """Baris berlabel 'Gambar 2: …' / '图 3 …' bila ada — keterangan gambar
    paling tepercaya karena penulis dokumen sendiri yang menuliskannya."""
    for ln in (teks or "").splitlines():
        m = _FIGUR_RE.match(ln.strip())
        if m and len(m.group("isi").strip()) >= 3:
            return m.group("isi").strip()
    return ""


def _caption(jalur: list[str], sekitar: str, teks_halaman: str) -> str:
    """Keterangan gambar dari TEKS DI SEKITARNYA — bukan hasil membaca gambar
    (tak ada kapabilitas vision). Dipakai untuk MENEMUKAN gambar sesuai konteks
    pertanyaan, jadi kualitasnya menentukan relevansi."""
    bagian: list[str] = []
    if jalur:
        bagian.append(jalur[-1])
    fig = _baris_figur(sekitar) or _baris_figur(teks_halaman)
    if fig:
        bagian.append(fig)
    if not fig:
        kal = re.split(r"(?<=[.。])\s+", (teks_halaman or "").strip())
        awal = " ".join(k for k in kal[:2] if k).strip()
        if awal:
            bagian.append(awal[:160])
    return " — ".join(x for x in bagian if x)[:MAX_CAPTION]


def _heading_pdf(teks: str, tumpukan: list[str]) -> list[str]:
    """Perbarui tumpukan bab dari baris-baris satu halaman (fallback bila PDF
    tak punya outline). Hanya mengisi breadcrumb — tak pernah mengubah
    pemotongan teks, jadi salah tebak berdampak kosmetik saja."""
    for ln in (teks or "").splitlines()[:40]:
        s = ln.strip()
        if not s or len(s) > 90 or s.endswith("."):
            continue
        m = _HEAD_NOMOR_RE.match(s)
        if m:
            lvl = m.group(1).count(".") + 1
            tumpukan[:] = tumpukan[:lvl - 1] + [s]
            continue
        m = _HEAD_CN_RE.match(s)
        if m:
            lvl = 1 if "章" in m.group(1) else 2
            tumpukan[:] = tumpukan[:lvl - 1] + [s]
    return list(tumpukan)


def _toc_pdf(pdf) -> list[tuple[int, int, str]]:
    """Outline/bookmark PDF → [(halaman_0based, level, judul)]. Eksak untuk
    dokumen terbitan (manual pabrikan, kebijakan) dan hampir gratis."""
    out = []
    try:
        for it in pdf.get_toc():
            try:
                out.append((int(it.page_index or 0), int(it.level or 0),
                            (it.title or "").strip()))
            except Exception:
                continue
    except Exception:
        return []
    return sorted(out, key=lambda t: t[0])


def _jalur_dari_toc(toc, hal_idx: int) -> list[str]:
    tumpukan: list[str] = []
    for hal, lvl, judul in toc:
        if hal > hal_idx:
            break
        if judul:
            tumpukan = tumpukan[:lvl] + [judul]
    return tumpukan


def _pil_ke_png(im, maks_sisi: int = 1600) -> bytes:
    """PIL → PNG teroptimasi tanpa encode/decode bolak-balik."""
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    if max(im.size) > maks_sisi:
        rasio = maks_sisi / max(im.size)
        im = im.resize((max(int(im.width * rasio), 1), max(int(im.height * rasio), 1)))
    buf = io.BytesIO()
    im.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def _gambar_halaman_pdf(pdfium, page, textpage, jalur, teks_halaman, lebar, tinggi):
    """Gambar embedded satu halaman PDF → [(sha1, png, caption)].

    Inilah yang HILANG di V1: dulu gambar hanya diambil dari halaman TANPA
    lapisan teks, sehingga semua diagram di dalam dokumen ber-teks tak pernah
    terindeks. Kegagalan di sini tak boleh menjatuhkan indexing teks.
    """
    out = []
    try:
        objek = list(page.get_objects(max_depth=2))
    except Exception:
        return out
    luas_hal = max(float(lebar) * float(tinggi), 1.0)
    for obj in objek:
        if len(out) >= MAX_GAMBAR_HALAMAN:
            break
        try:
            if getattr(obj, "type", None) != pdfium.raw.FPDF_PAGEOBJ_IMAGE:
                continue
            w, h = obj.get_px_size()
            if w < MIN_SISI_GAMBAR or h < MIN_SISI_GAMBAR:
                continue
            if w * h > MAX_PX_GAMBAR:      # cek SEBELUM decode — kunci RAM
                continue
            # get_bounds() → (left, bottom, right, top) dalam titik halaman.
            l, b, r, t = obj.get_bounds()
            if (abs(r - l) * abs(t - b)) / luas_hal < 0.015:
                continue                   # terlalu kecil relatif halaman
            if b >= tinggi * 0.92 or t <= tinggi * 0.08:
                continue                   # pita kop/footer → logo
            try:
                bitmap = obj.get_bitmap(render=False)
            except Exception:
                bitmap = obj.get_bitmap(render=True)
            try:
                png = _pil_ke_png(bitmap.to_pil())
            finally:
                try:
                    bitmap.close()
                except Exception:
                    pass
            sekitar = ""
            try:
                sekitar = (textpage.get_text_bounded(
                    left=l - 20, bottom=b - 90, right=r + 20, top=b) or "")
                if not sekitar.strip():
                    sekitar = (textpage.get_text_bounded(
                        left=l - 20, bottom=t, right=r + 20, top=t + 60) or "")
            except Exception:
                sekitar = ""
            out.append((hashlib.sha1(png).hexdigest(), png,
                        _caption(jalur, sekitar, teks_halaman)))
        except Exception:
            continue           # satu objek rusak tak menjatuhkan halaman
    return out


def dari_pdf(data: bytes) -> list[dict]:
    """Teks + tabel + gambar + breadcrumb per halaman. Halaman ditutup tiap
    iterasi — itu kunci supaya PDF 200 halaman tidak menumpuk di RAM."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        raise ValueError("Dukungan PDF belum terpasang di server (pypdfium2).")
    try:
        pdf = pdfium.PdfDocument(io.BytesIO(data))
    except Exception:
        raise ValueError("Berkas PDF rusak / terkunci sandi.")
    # Probe sekali: bila API objek pypdfium2 berbeda di versi terpasang, fitur
    # gambar dimatikan dan indexing TEKS tetap jalan seperti biasa. Nama metode
    # ini pernah berbeda antar versi (get_pos vs get_bounds) — jangan diasumsi.
    bisa_gambar = (hasattr(pdfium, "raw")
                   and hasattr(pdfium.raw, "FPDF_PAGEOBJ_IMAGE")
                   and hasattr(pdfium.PdfImage, "get_bounds")
                   and hasattr(pdfium.PdfImage, "get_px_size"))

    out: list[dict] = []
    n_gambar = 0
    byte_gambar = 0
    kandidat: list[tuple[int, str, bytes, str]] = []   # (idx_bagian, sha1, png, caption)
    hash_halaman: dict[str, set[int]] = {}
    tumpukan: list[str] = []
    try:
        toc = _toc_pdf(pdf)
        total = min(len(pdf), MAX_HALAMAN_PDF)
        for i in range(total):
            page = None
            try:
                page = pdf[i]
                textpage = None
                try:
                    textpage = page.get_textpage()
                    teks = textpage.get_text_bounded() or ""
                except Exception:
                    teks = ""
                jalur = (_jalur_dari_toc(toc, i) if toc
                         else _heading_pdf(teks, tumpukan))

                tabel = _baris_tabel_pdf(teks)
                if tabel:
                    out.extend(_tabel_jadi_bagian(tabel, i + 1, "", jalur))
                idx_teks_pertama = len(out)
                for t in potong_teks(teks):
                    out.append(_bagian(teks=t, halaman=i + 1, jalur=jalur))
                ada_teks = len(out) > idx_teks_pertama

                if not teks.strip() and n_gambar < MAX_GAMBAR and byte_gambar < MAX_GAMBAR_BYTES:
                    # Halaman tanpa lapisan teks (pindaian) → render @150 DPI.
                    try:
                        png = _pil_ke_png(page.render(scale=150 / 72).to_pil())
                        cap = _caption(jalur, "", "")
                        b = _bagian(teks=cap, halaman=i + 1, jalur=jalur,
                                    gambar=[png],
                                    gambar_meta=[{"caption": cap, "halaman": i + 1}])
                        out.append(b)
                        n_gambar += 1
                        byte_gambar += len(png)
                    except Exception:
                        pass
                elif bisa_gambar and textpage is not None and n_gambar < MAX_GAMBAR \
                        and byte_gambar < MAX_GAMBAR_BYTES:
                    try:
                        w_hal, h_hal = page.get_size()
                    except Exception:
                        w_hal, h_hal = 595.0, 842.0
                    for sha, png, cap in _gambar_halaman_pdf(
                            pdfium, page, textpage, jalur, teks, w_hal, h_hal):
                        if n_gambar >= MAX_GAMBAR or byte_gambar >= MAX_GAMBAR_BYTES:
                            break
                        hash_halaman.setdefault(sha, set()).add(i)
                        # Ditempelkan ke chunk TEKS halaman ini — chunk gambar
                        # tanpa teks skornya selalu 0 alias tak pernah ketemu.
                        target = idx_teks_pertama if ada_teks else -1
                        kandidat.append((target, sha, png, cap))
                        n_gambar += 1
                        byte_gambar += len(png)
            except Exception:
                continue    # satu halaman rusak tak menjatuhkan sisanya
            finally:
                if page is not None:
                    try:
                        page.close()
                    except Exception:
                        pass
    finally:
        try:
            pdf.close()
        except Exception:
            pass

    # Saring logo/watermark: hash yang sama muncul di banyak halaman berbeda.
    # Dokumen pendek dikecualikan — wajar mengulang satu diagram di sana.
    saring = total >= 5
    dipakai: set[str] = set()
    for target, sha, png, cap in kandidat:
        if saring and len(hash_halaman.get(sha, ())) >= LOGO_MIN_HALAMAN:
            continue
        if sha in dipakai:
            continue
        dipakai.add(sha)
        if 0 <= target < len(out):
            out[target]["gambar"].append(png)
            out[target]["gambar_meta"].append({"caption": cap,
                                               "halaman": out[target]["halaman"]})
    return out


def pdf_tanpa_teks(bagian: list[dict]) -> bool:
    """True bila PDF sama sekali tak punya lapisan teks (hasil pindaian) —
    pemanggil memberi tahu admin agar menambah keterangan manual."""
    return bool(bagian) and not any((b.get("teks") or "").strip() for b in bagian)


# ── DOCX (lazy import python-docx, fallback stdlib) ──────────────────
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _docx_stdlib(data: bytes) -> list[dict]:
    """Fallback tanpa dependency: .docx = ZIP, ambil teks dari document.xml.
    Kualitas lebih rendah (tanpa tabel/heading) tapi fitur tetap jalan."""
    from xml.etree import ElementTree
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml = z.read("word/document.xml")
    except Exception:
        raise ValueError("Berkas Word rusak / bukan .docx yang valid.")
    root = ElementTree.fromstring(xml)
    paragraf = []
    for p in root.iter(f"{_W_NS}p"):
        teks = "".join(t.text or "" for t in p.iter(f"{_W_NS}t"))
        if teks.strip():
            paragraf.append(teks.strip())
    return [_bagian(teks=t) for t in potong_teks("\n".join(paragraf))]


def dari_docx(data: bytes) -> list[dict]:
    _cek_zip(data)
    try:
        import docx  # type: ignore
    except ImportError:
        return _docx_stdlib(data)
    try:
        d = docx.Document(io.BytesIO(data))
    except Exception:
        raise ValueError("Berkas Word rusak / bukan .docx yang valid.")
    out: list[dict] = []
    tumpukan: list[str] = []      # breadcrumb bab→subbab dari gaya Heading N
    buf: list[str] = []
    tunda: list[tuple[bytes, str]] = []   # gambar inline menunggu bagian terdekat
    dipakai_rid: set[str] = set()
    n_gambar = 0

    def _flush():
        if not buf:
            return
        jalur = list(tumpukan)
        potong = potong_teks("\n".join(buf))
        buf.clear()
        awal = len(out)
        for t in potong:
            out.append(_bagian(teks=t, sub=(jalur[-1] if jalur else ""), jalur=jalur))
        # Gambar yang ditemui di paragraf ini menempel ke bagian pertamanya —
        # dengan begitu ia mewarisi breadcrumb + teks di sekitarnya.
        if tunda and len(out) > awal:
            for png, cap in tunda:
                out[awal]["gambar"].append(png)
                out[awal]["gambar_meta"].append({"caption": cap, "halaman": 0})
        tunda.clear()

    for p in d.paragraphs:
        teks = (p.text or "").strip()
        gaya = (getattr(p.style, "name", "") or "")
        # Gambar inline, dalam urutan dokumen (V1 memindainya di AKHIR sehingga
        # kehilangan posisi & konteks).
        try:
            for blip in p._p.iter(f"{_A_NS}blip"):
                if n_gambar >= MAX_GAMBAR:
                    break
                rid = blip.get(f"{_R_NS}embed")
                if not rid or rid in dipakai_rid:
                    continue
                dipakai_rid.add(rid)
                try:
                    png = optimalkan_gambar(d.part.related_parts[rid].blob)
                except Exception:
                    continue
                tunda.append((png, _caption(tumpukan, teks, teks)))
                n_gambar += 1
        except Exception:
            pass
        if gaya.startswith("Heading") and teks:
            _flush()
            try:
                lvl = max(int(gaya.split()[-1]), 1)
            except Exception:
                lvl = 1
            tumpukan[:] = tumpukan[:lvl - 1] + [teks]
            buf.append(teks)
            continue
        if teks:
            buf.append(teks)
    _flush()

    jalur_akhir = list(tumpukan)
    for t in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        out.extend(_tabel_jadi_bagian(
            rows, 0, (jalur_akhir[-1] if jalur_akhir else ""), jalur_akhir))

    # Fallback gambar mengambang (floating shape/VML) yang tak punya blip di
    # paragraf mana pun — dilekatkan ke bagian TERAKHIR yang punya teks, bukan
    # dibiarkan jadi bagian yatim tanpa teks (yang skornya selalu 0).
    sisa = [b for b in out if (b.get("teks") or "").strip()]
    for rel in d.part.related_parts.values():
        if n_gambar >= MAX_GAMBAR or not sisa:
            break
        try:
            if "image" not in (getattr(rel, "content_type", "") or ""):
                continue
            if getattr(rel, "partname", None) and str(rel.partname) in dipakai_rid:
                continue
            blob = rel.blob
            if any(hashlib.sha1(blob).hexdigest() == hashlib.sha1(g).hexdigest()
                   for b in out for g in b.get("gambar", [])):
                continue
            png = optimalkan_gambar(blob)
            sisa[-1]["gambar"].append(png)
            sisa[-1]["gambar_meta"].append(
                {"caption": _caption(jalur_akhir, "", sisa[-1].get("teks", "")),
                 "halaman": 0})
            n_gambar += 1
        except Exception:
            continue
    return out


# ── dispatcher ───────────────────────────────────────────────────────
def ekstrak(data: bytes, nama: str) -> list[dict]:
    """Berkas → daftar bagian. Melempar ValueError berpesan Indonesia bila
    formatnya tak didukung atau berkasnya rusak."""
    ext = Path(nama or "").suffix.lower()
    if ext == ".txt":
        return dari_txt(data)
    if ext == ".csv":
        return dari_csv(data)
    if ext in (".xlsx", ".xlsm"):
        return dari_excel(data)
    if ext == ".pdf":
        return dari_pdf(data)
    if ext == ".docx":
        return dari_docx(data)
    if ext in (".png", ".jpg", ".jpeg"):
        return dari_gambar(data)
    raise ValueError(f"Format '{ext or nama}' belum didukung.")
