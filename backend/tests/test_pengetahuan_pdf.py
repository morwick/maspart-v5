"""Ekstraksi PDF & DOCX untuk indexing pengetahuan.

Fixture PDF dibuat pakai reportlab (sudah ada di requirements runtime — tak
perlu menambah dependency test), DOCX pakai python-docx.
"""
import io
import zipfile

import pytest

from app.services import knowledge_util, pengetahuan
from app.services import pengetahuan_extract as ext
from app.services import pengetahuan_index as idx


@pytest.fixture(autouse=True)
def _tmp_store(tmp_path, monkeypatch):
    monkeypatch.setattr(pengetahuan, "_dir", lambda: tmp_path)
    monkeypatch.setattr(idx, "get_settings",
                        lambda: type("S", (), {"ai_configured": False})())
    knowledge_util._LOAD_CACHE.clear()
    yield tmp_path
    knowledge_util._LOAD_CACHE.clear()


def _pdf(halaman: list[list[str]]) -> bytes:
    """PDF sederhana: tiap elemen = daftar baris untuk satu halaman."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    for baris in halaman:
        y = 800
        for ln in baris:
            c.drawString(60, y, ln)
            y -= 18
        c.showPage()
    c.save()
    return buf.getvalue()


def _pdf_bergambar(baris_teks, warna=(200, 30, 30), sisi=300, halaman=1,
                   logo=False):
    """PDF ber-TEKS yang juga memuat gambar — kasus yang di V1 gambarnya
    hilang total karena hanya halaman tanpa teks yang diproses."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas
    from PIL import Image
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    img = ImageReader(Image.new("RGB", (sisi, sisi), warna))
    kecil = ImageReader(Image.new("RGB", (24, 24), (9, 9, 9)))
    for n in range(halaman):
        y = 780
        for ln in baris_teks:
            c.drawString(60, y, ln)
            y -= 18
        c.drawImage(img, 80, 320, width=260, height=260)
        if logo:                       # gambar identik di tiap halaman = logo
            c.drawImage(kecil, 60, 20, width=24, height=24)
        c.showPage()
    c.save()
    return buf.getvalue()


def _docx(paragraf=(), tabel=None, heading=""):
    import docx
    d = docx.Document()
    if heading:
        d.add_heading(heading, level=1)
    for p in paragraf:
        d.add_paragraph(p)
    if tabel:
        t = d.add_table(rows=len(tabel), cols=len(tabel[0]))
        for i, baris in enumerate(tabel):
            for j, sel in enumerate(baris):
                t.cell(i, j).text = str(sel)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── PDF ──────────────────────────────────────────────────────────────
def test_pdf_teks_dan_nomor_halaman_benar():
    data = _pdf([["Prosedur retur barang"], ["Syarat klaim garansi"]])
    bagian = ext.ekstrak(data, "kebijakan.pdf")
    teks = {b["halaman"]: b["teks"] for b in bagian if b["teks"]}
    assert "retur" in teks[1].lower()
    assert "garansi" in teks[2].lower()


def test_pdf_rusak_melempar_pesan_indonesia():
    with pytest.raises(ValueError, match="rusak"):
        ext.ekstrak(b"%PDF-1.7 tapi isinya sampah", "rusak.pdf")


def test_pdf_batas_halaman_dihormati(monkeypatch):
    monkeypatch.setattr(ext, "MAX_HALAMAN_PDF", 2)
    data = _pdf([[f"halaman {i}"] for i in range(6)])
    bagian = ext.ekstrak(data, "tebal.pdf")
    assert max(b["halaman"] for b in bagian) <= 2


def test_pdf_tanpa_lapisan_teks_dirender_jadi_gambar():
    """Halaman kosong = tak ada lapisan teks (mirip hasil pindaian) → gambar."""
    data = _pdf([[]])
    bagian = ext.ekstrak(data, "pindaian.pdf")
    assert ext.pdf_tanpa_teks(bagian)
    assert any(b["gambar"] for b in bagian)


def test_pdf_pindaian_memberi_peringatan_ke_admin(_tmp_store):
    dok = pengetahuan.add_dokumen("Hasil Pindai", teks_admin="x")
    d = pengetahuan.berkas_dir()
    d.mkdir(parents=True, exist_ok=True)
    (d / f"{dok['id']}_0.pdf").write_bytes(_pdf([[]]))
    pengetahuan.update_dokumen(dok["id"], teks_admin="", berkas=[
        {"nama": "pindaian.pdf", "nama_simpan": f"{dok['id']}_0.pdf", "ext": ".pdf"}])
    knowledge_util._LOAD_CACHE.clear()
    idx.proses(dok["id"])
    doc = pengetahuan.get_dokumen(dok["id"])
    assert doc["status"] == "selesai_sebagian"
    assert "pindaian" in doc["error"].lower()


def test_tabel_pdf_direkonstruksi_dari_whitespace():
    baris = ["Part      Torsi     Satuan",
             "baut roda      600       Nm",
             "baut as        400       Nm"]
    tabel = ext._baris_tabel_pdf("\n".join(baris))
    assert tabel and tabel[0][0].startswith("Part")
    assert any("baut roda" in r[0] for r in tabel)


def test_teks_biasa_tidak_disalahartikan_sebagai_tabel():
    prosa = "Retur barang diajukan maksimal tujuh hari setelah barang diterima."
    assert ext._baris_tabel_pdf(prosa) == []


# ── DOCX ─────────────────────────────────────────────────────────────
def test_docx_paragraf_heading_dan_tabel():
    data = _docx(paragraf=["Retur diajukan maksimal 7 hari."],
                 tabel=[["Part", "Torsi"], ["baut roda", "600"]],
                 heading="Kebijakan Retur")
    bagian = ext.ekstrak(data, "kebijakan.docx")
    teks = " ".join(b["teks"] for b in bagian)
    assert "Retur diajukan" in teks
    assert any(b["sub"] == "Kebijakan Retur" for b in bagian)
    assert any(b["tabel"] and b["tabel"][0] == ["Part", "Torsi"] for b in bagian)


def test_docx_rusak_melempar_pesan_indonesia():
    with pytest.raises(ValueError):
        ext.ekstrak(b"bukan zip", "rusak.docx")


def test_docx_fallback_stdlib_saat_python_docx_absen(monkeypatch):
    """Fitur harus tetap jalan bila dependency python-docx tak terpasang."""
    data = _docx(paragraf=["Retur diajukan maksimal 7 hari."])
    asli = __import__

    def tanpa_docx(nama, *a, **kw):
        if nama == "docx":
            raise ImportError("tidak terpasang")
        return asli(nama, *a, **kw)
    monkeypatch.setattr("builtins.__import__", tanpa_docx)
    bagian = ext.ekstrak(data, "kebijakan.docx")
    assert "Retur diajukan" in " ".join(b["teks"] for b in bagian)


def test_docx_zip_bomb_ditolak(monkeypatch):
    monkeypatch.setattr(ext, "MAX_ZIP_URAI", 100)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/document.xml", "x" * 5000)
    with pytest.raises(ValueError, match="mengembang"):
        ext.ekstrak(buf.getvalue(), "bom.docx")


# ── gambar embedded di PDF ber-teks (celah utama V1) ─────────────────
def test_pdf_berteks_mengekstrak_gambar_embedded():
    data = _pdf_bergambar(["Prosedur retur barang", "Gambar 1: alur pengajuan retur"])
    bagian = ext.ekstrak(data, "kebijakan.pdf")
    ber = [b for b in bagian if b["gambar"]]
    assert ber, "gambar embedded di halaman ber-teks tidak terekstrak"
    assert (ber[0]["teks"] or "").strip(), "gambar harus menempel ke chunk BERTEKS"


def test_caption_gambar_dari_baris_figur():
    data = _pdf_bergambar(["Prosedur retur", "Gambar 1: alur pengajuan retur"])
    bagian = ext.ekstrak(data, "kebijakan.pdf")
    caps = [m.get("caption", "") for b in bagian for m in b.get("gambar_meta", [])]
    assert any("alur pengajuan retur" in c for c in caps), caps


def test_gambar_kecil_dan_logo_berulang_dibuang():
    data = _pdf_bergambar(["Prosedur retur barang dagangan"], halaman=6, logo=True)
    bagian = ext.ekstrak(data, "kebijakan.pdf")
    # logo 24px kena filter ukuran; gambar utama identik di 6 halaman kena
    # filter "hash sama di >=3 halaman" → hanya menyisakan sedikit/nol
    total = sum(len(b["gambar"]) for b in bagian)
    assert total <= 1, f"logo/watermark tidak tersaring (total {total})"


def test_ekstraksi_gambar_gagal_tidak_menjatuhkan_indexing(monkeypatch):
    """Kegagalan API gambar tak boleh membuat isi TEKS ikut hilang."""
    def boom(*a, **kw):
        raise RuntimeError("API berubah")
    monkeypatch.setattr(ext, "_gambar_halaman_pdf", boom)
    bagian = ext.ekstrak(_pdf_bergambar(["Prosedur retur barang"]), "a.pdf")
    assert any("retur" in (b["teks"] or "").lower() for b in bagian)


def test_breadcrumb_pdf_heuristik_bernomor():
    data = _pdf(["3. Retur", "3.2 Syarat pengajuan",
                 "Barang wajib utuh dan disertai foto kondisi."])
    bagian = ext.ekstrak(data, "kebijakan.pdf")
    jalur = [b.get("jalur") for b in bagian if b.get("jalur")]
    assert jalur and any("3.2" in " ".join(j) for j in jalur), jalur


def test_breadcrumb_docx_bertingkat():
    import docx
    d = docx.Document()
    d.add_heading("3 Retur", level=1)
    d.add_heading("3.2 Syarat", level=2)
    d.add_paragraph("Barang wajib utuh dan disertai foto kondisi saat diterima.")
    buf = io.BytesIO()
    d.save(buf)
    bagian = ext.ekstrak(buf.getvalue(), "kebijakan.docx")
    isi = [b for b in bagian if "wajib utuh" in (b["teks"] or "")]
    assert isi and isi[0]["jalur"] == ["3 Retur", "3.2 Syarat"]


def test_docx_gambar_mewarisi_heading_dan_menempel_ke_teks():
    import docx
    from PIL import Image
    img = io.BytesIO()
    Image.new("RGB", (200, 150), (10, 120, 200)).save(img, format="PNG")
    img.seek(0)
    d = docx.Document()
    d.add_heading("4 Pemasangan", level=1)
    p = d.add_paragraph("Langkah pemasangan filter oli mesin.")
    p.add_run().add_picture(img)
    buf = io.BytesIO()
    d.save(buf)
    bagian = ext.ekstrak(buf.getvalue(), "panduan.docx")
    ber = [b for b in bagian if b["gambar"]]
    assert ber, "gambar inline DOCX tak terambil"
    assert ber[0]["jalur"] == ["4 Pemasangan"]
    assert (ber[0]["teks"] or "").strip()


def test_excel_gambar_terambil():
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from PIL import Image
    p = io.BytesIO()
    Image.new("RGB", (180, 140), (30, 160, 90)).save(p, format="PNG")
    p.seek(0)
    wb = Workbook()
    ws = wb.active
    ws.title = "Torsi"
    ws.append(["Part", "Nm"])
    ws.append(["baut roda", 600])
    ws.add_image(XLImage(p), "D2")
    buf = io.BytesIO()
    wb.save(buf)
    bagian = ext.ekstrak(buf.getvalue(), "torsi.xlsx")
    assert any(b["gambar"] for b in bagian), "gambar Excel tak terambil"


def test_kolom_tabel_tersimpan():
    data = b"Part;Torsi;Satuan\nbaut roda;600;Nm\n"
    bagian = ext.ekstrak(data, "torsi.csv")
    assert bagian[0]["kolom"] == ["Part", "Torsi", "Satuan"]


# ── end-to-end lewat job ─────────────────────────────────────────────
def test_pdf_terindeks_dan_bisa_dicari(_tmp_store):
    dok = pengetahuan.add_dokumen("Kebijakan", teks_admin="x")
    d = pengetahuan.berkas_dir()
    d.mkdir(parents=True, exist_ok=True)
    (d / f"{dok['id']}_0.pdf").write_bytes(
        _pdf([["Prosedur retur barang dagangan"]]))
    pengetahuan.update_dokumen(dok["id"], teks_admin="", berkas=[
        {"nama": "kebijakan.pdf", "nama_simpan": f"{dok['id']}_0.pdf", "ext": ".pdf"}])
    knowledge_util._LOAD_CACHE.clear()
    idx.proses(dok["id"])
    knowledge_util._LOAD_CACHE.clear()
    hasil = pengetahuan.search("retur")
    assert hasil and hasil[0]["sumber"] == "kebijakan.pdf"
    assert hasil[0]["tipe"] == "pdf" and hasil[0]["halaman"] == 1
